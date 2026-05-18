import docx
import sys
import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

path = r'C:\Users\Eduardo DELL\OneDrive\Área de Trabalho\vanguard\CLIENTES\VALDECE\Contrato_Toga_Digital_Valdece_19052026.docx'
out  = r'C:\Users\Eduardo DELL\OneDrive\Área de Trabalho\vanguard\CLIENTES\VALDECE\Contrato_Toga_Digital_Valdece_19052026_v2.docx'

doc = docx.Document(path)

# Localizar indice do paragrafo "4. VALOR E PAGAMENTO"
idx_clausula4 = None
for i, p in enumerate(doc.paragraphs):
    if '4. VALOR E PAGAMENTO' in p.text or '4.VALOR' in p.text.upper():
        idx_clausula4 = i
        break

if idx_clausula4 is None:
    print("ERRO: Clausula 4 nao encontrada.")
    sys.exit(1)

print(f"Clausula 4 encontrada no paragrafo [{idx_clausula4}]: {doc.paragraphs[idx_clausula4].text}")
print(f"Proximo paragrafo [{idx_clausula4+1}]: '{doc.paragraphs[idx_clausula4+1].text}'")
print(f"Proximo paragrafo [{idx_clausula4+2}]: '{doc.paragraphs[idx_clausula4+2].text}'")

# Conteudo da Clausula 4 corrigida
linhas_clausula4 = [
    "4.1  Valor fixo: R$ 5.000,00 (cinco mil reais) — pagamento único, a ser realizado conforme acordo entre as partes.",
    "",
    "4.2  Participação em Resultados (Revenue Share): Sobre qualquer receita recorrente gerada por produto SaaS derivado desta entrega, o Contratante repassará ao Contratado 20% (vinte por cento) do MRR (Monthly Recurring Revenue) apurado mensalmente, a partir do primeiro mês de operação comercial. O repasse será realizado até o dia 10 (dez) de cada mês subsequente, mediante comprovante de faturamento.",
    "",
    "4.3  O valor fixo estipulado no item 4.1 foi estabelecido abaixo do valor de mercado exclusivamente em função da contrapartida de participação em resultados prevista no item 4.2. A ausência de cumprimento do item 4.2 implica reequilíbrio contratual a ser negociado entre as partes.",
]

# Obter o elemento XML do paragrafo da clausula 4 para inserir depois dele
p_clausula4 = doc.paragraphs[idx_clausula4]._element
parent = p_clausula4.getparent()

# Criar paragrafos novos e inserir apos o titulo da clausula 4
# Vamos inserir antes do paragrafo seguinte (que e o vazio ou ja o 5)
p_next = doc.paragraphs[idx_clausula4 + 1]._element

def novo_paragrafo(doc, texto, bold=False):
    p = OxmlElement('w:p')
    if texto:
        r = OxmlElement('w:r')
        if bold:
            rpr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rpr.append(b)
            r.append(rpr)
        t = OxmlElement('w:t')
        t.text = texto
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
    return p

# Inserir os paragrafos da clausula 4 apos o titulo
ref = p_clausula4
for linha in linhas_clausula4:
    novo_p = novo_paragrafo(doc, linha)
    ref.addnext(novo_p)
    ref = novo_p

doc.save(out)
print(f"\n[OK] Contrato salvo em: {out}")
print("Clausula 4 inserida com 4.1 (valor fixo) + 4.2 (Revenue Share 20% MRR) + 4.3 (fundamento do desconto).")
